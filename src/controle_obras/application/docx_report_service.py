"""Gerador de relatório DOCX com conversão para PDF via LibreOffice."""

import logging
import os
import subprocess
import tempfile
from datetime import datetime
from decimal import Decimal
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


logger = logging.getLogger(__name__)


def _formatar_data(valor: Any) -> str:
    """Formata data para dd/mm/aaaa."""
    if valor is None:
        return ""
    if hasattr(valor, "strftime"):
        return valor.strftime("%d/%m/%Y")
    return str(valor)


def _formatar_moeda(valor: Decimal) -> str:
    """Formata valor monetário para R$ 1.234,56."""
    fmt = f"{float(valor):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return f"R$ {fmt}"


def _texto(valor: Any, padrao: str = "") -> str:
    """Retorna string segura, tratando None."""
    return padrao if valor is None else str(valor)


class DocxReportService:
    """Serviço de geração de relatórios DOCX/PDF."""

    def __init__(
        self,
        obra_service: Any,
        aditivo_service: Any,
        lancamento_service: Any,
        anexo_service: Any,
        resumo_service: Any,
        relatorio_repo: Any,
        storage: Any,
        empresa_service: Any | None = None,
    ) -> None:
        self._obra_service = obra_service
        self._aditivo_service = aditivo_service
        self._lancamento_service = lancamento_service
        self._anexo_service = anexo_service
        self._resumo_service = resumo_service
        self._relatorio_repo = relatorio_repo
        self._storage = storage
        self._empresa_service = empresa_service

        # Caminhos do LibreOffice no Windows
        self._libreoffice_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.com",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.com",
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]

    def _obter_libreoffice_path(self) -> str | None:
        """Localiza o executável do LibreOffice."""
        for path in self._libreoffice_paths:
            if os.path.exists(path):
                return path
        return None

    def _obter_responsavel(self) -> str:
        """Obtém o nome do responsável da empresa."""
        if self._empresa_service is None:
            return ""
        try:
            empresa = self._empresa_service.obter()
            if empresa and empresa.responsavel:
                return empresa.responsavel
            return ""
        except Exception:
            return ""

    def _obter_cnpj(self) -> str:
        """Obtém o CNPJ da empresa."""
        if self._empresa_service is None:
            return ""
        try:
            empresa = self._empresa_service.obter()
            if empresa and empresa.cnpj:
                return empresa.cnpj
            return ""
        except Exception:
            return ""

    def _formatar_tamanho(self, tamanho_bytes: int) -> str:
        """Formata tamanho de arquivo."""
        if tamanho_bytes >= 1024 * 1024:
            return f"{tamanho_bytes / (1024 * 1024):.1f} MB"
        elif tamanho_bytes >= 1024:
            return f"{tamanho_bytes / 1024:.1f} KB"
        return f"{tamanho_bytes} B"

    def _criar_documento(self) -> Document:
        """Cria um documento DOCX com configuração A4 e margens."""
        doc = Document()
        
        # Configurar página A4
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        
        return doc

    def _adicionar_titulo(self, doc: Document, texto: str, tamanho: int = 14, centralizado: bool = True, sublinhado: bool = False) -> None:
        """Adiciona um título ao documento."""
        p = doc.add_paragraph()
        if centralizado:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(texto)
        run.font.name = "Arial"
        run.font.size = Pt(tamanho)
        run.font.bold = True
        if sublinhado:
            run.underline = True
        doc.add_paragraph()

    def _adicionar_secao(self, doc: Document, titulo: str) -> None:
        """Adiciona título de seção."""
        p = doc.add_paragraph()
        run = p.add_run(titulo)
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.font.bold = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

    def _adicionar_tabela_resumo(self, doc: Document, resumo: Any) -> None:
        """Adiciona tabela de resumo financeiro."""
        table = doc.add_table(rows=2, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Configurar larguras iguais
        for col in table.columns:
            col.width = Cm(4.5)
        
        # Cabeçalho
        cabecalhos = ["Valor Contratado", "Total Aditivos", "Total Gasto", "Valor Líquido"]
        for i, texto in enumerate(cabecalhos):
            celula = table.cell(0, i)
            celula.text = texto
            for paragraph in celula.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
                    run.font.bold = True
        
        # Valores
        valores = [
            _formatar_moeda(resumo.valor_contratado),
            _formatar_moeda(resumo.total_aditivos),
            _formatar_moeda(resumo.total_gasto),
            _formatar_moeda(resumo.valor_liquido),
        ]
        for i, valor in enumerate(valores):
            celula = table.cell(1, i)
            celula.text = valor
            for paragraph in celula.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(11)
                    run.font.bold = True

    def _adicionar_tabela_aditivos(self, doc: Document, aditivos: list) -> None:
        """Adiciona tabela de aditivos."""
        table = doc.add_table(rows=len(aditivos) + 1, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Configurar larguras
        table.columns[0].width = Cm(3.0)   # 20%
        table.columns[1].width = Cm(8.25)  # 55%
        table.columns[2].width = Cm(3.75)  # 25%
        
        # Cabeçalho
        cabecalhos = ["Data", "Descrição", "Valor"]
        for i, texto in enumerate(cabecalhos):
            celula = table.cell(0, i)
            celula.text = texto
            for paragraph in celula.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if i < 2 else WD_ALIGN_PARAGRAPH.RIGHT
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
                    run.font.bold = True
        
        # Dados
        for row_idx, aditivo in enumerate(aditivos, start=1):
            data = _formatar_data(aditivo.data_aditivo)
            descricao = _texto(aditivo.descricao, "Sem descrição")
            valor = _formatar_moeda(aditivo.valor)
            
            table.cell(row_idx, 0).text = data
            table.cell(row_idx, 1).text = descricao
            table.cell(row_idx, 2).text = valor
            
            for col_idx in range(3):
                celula = table.cell(row_idx, col_idx)
                for paragraph in celula.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx < 2 else WD_ALIGN_PARAGRAPH.RIGHT
                    for run in paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(9)

    def _adicionar_tabela_lancamentos(self, doc: Document, lancamentos: list) -> None:
        """Adiciona tabela de lançamentos."""
        table = doc.add_table(rows=len(lancamentos) + 1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Configurar larguras
        table.columns[0].width = Cm(2.25)  # 15%
        table.columns[1].width = Cm(6.0)   # 40%
        table.columns[2].width = Cm(3.0)   # 20%
        table.columns[3].width = Cm(3.75)  # 25%
        
        # Cabeçalho
        cabecalhos = ["Data", "Descrição", "Tipo", "Valor"]
        for i, texto in enumerate(cabecalhos):
            celula = table.cell(0, i)
            celula.text = texto
            for paragraph in celula.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
                    run.font.bold = True
        
        # Dados
        for row_idx, lanc in enumerate(lancamentos, start=1):
            data = _formatar_data(lanc.data_lancamento)
            descricao = _texto(lanc.descricao, "Sem descrição")
            tipo = _texto(lanc.tipo_nome, "Não informado") if hasattr(lanc, "tipo_nome") else "Não informado"
            valor = _formatar_moeda(lanc.valor_total)
            
            table.cell(row_idx, 0).text = data
            table.cell(row_idx, 1).text = descricao
            table.cell(row_idx, 2).text = tipo
            table.cell(row_idx, 3).text = valor
            
            for col_idx in range(4):
                celula = table.cell(row_idx, col_idx)
                for paragraph in celula.paragraphs:
                    if col_idx == 3:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(9)

    def _adicionar_anexos(self, doc: Document, anexos: list) -> None:
        """Adiciona lista de anexos."""
        for i, anexo in enumerate(anexos):
            nome = _texto(anexo.nome_original, "Sem nome")
            tipo = _texto(anexo.tipo_anexo, "Não informado")
            data_doc = anexo.data_documento or (anexo.created_at.date() if anexo.created_at else None)
            data = _formatar_data(data_doc)
            tamanho = self._formatar_tamanho(anexo.tamanho_bytes or 0)
            
            # Nome do arquivo (pode ocupar várias linhas)
            p_nome = doc.add_paragraph()
            run_nome = p_nome.add_run(nome)
            run_nome.font.name = "Arial"
            run_nome.font.size = Pt(9)
            
            # Metadados
            p_meta = doc.add_paragraph()
            meta_texto = f"Tipo: {tipo} | Data: {data} | Tamanho: {tamanho}"
            run_meta = p_meta.add_run(meta_texto)
            run_meta.font.name = "Arial"
            run_meta.font.size = Pt(8)
            
            # Linha separadora (exceto no último)
            if i < len(anexos) - 1:
                p_sep = doc.add_paragraph()
                p_sep.paragraph_format.space_before = Pt(2)
                p_sep.paragraph_format.space_after = Pt(2)
                run_sep = p_sep.add_run("─" * 80)
                run_sep.font.name = "Arial"
                run_sep.font.size = Pt(6)
                run_sep.font.color.rgb = None  # Cinza padrão

    def _adicionar_assinatura(self, doc: Document, responsavel: str, cnpj: str) -> None:
        """Adiciona seção de assinatura."""
        doc.add_paragraph()  # Espaçamento
        
        # Nome do responsável
        p_nome = doc.add_paragraph()
        p_nome.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_nome = p_nome.add_run(responsavel if responsavel else "Cassio Reis")
        run_nome.font.name = "Arial"
        run_nome.font.size = Pt(11)
        run_nome.font.bold = True
        
        # CNPJ ou cargo
        p_cargo = doc.add_paragraph()
        p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cnpj:
            run_cargo = p_cargo.add_run(f"CNPJ: {cnpj}")
        else:
            run_cargo = p_cargo.add_run("Responsável Legal")
        run_cargo.font.name = "Arial"
        run_cargo.font.size = Pt(9)

    def _converter_para_pdf(self, docx_path: Path) -> Path:
        """Converte DOCX para PDF usando LibreOffice."""
        libreoffice = self._obter_libreoffice_path()
        if not libreoffice:
            raise RuntimeError(
                "LibreOffice não foi encontrado. "
                "Instale em: https://www.libreoffice.org/download/"
            )
        
        output_dir = docx_path.parent
        output_pdf = docx_path.with_suffix(".pdf")
        
        # Remover PDF existente se houver
        if output_pdf.exists():
            output_pdf.unlink()
        
        cmd = [
            libreoffice,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", str(output_dir),
            str(docx_path),
        ]
        
        logger.info(f"Executando: {' '.join(cmd)}")
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        
        if result.returncode != 0:
            logger.error(f"LibreOffice stderr: {result.stderr}")
            raise RuntimeError(f"Erro ao converter DOCX para PDF: {result.stderr}")
        
        if not output_pdf.exists():
            raise RuntimeError("PDF não foi gerado após conversão")
        
        return output_pdf

    def gerar_relatorio_obra_docx(self, obra_id: int) -> Path:
        """Gera relatório PDF via DOCX + LibreOffice."""
        print(f"[DOCX] Gerador: python-docx + LibreOffice")
        print(f"[DOCX] Obra: {obra_id}")
        
        # Obter dados
        obra = self._obra_service.obter(obra_id)
        if not obra:
            raise ValueError(f"Obra {obra_id} não encontrada.")
        
        resumo = self._resumo_service.calcular_resumo(obra_id)
        aditivos = self._aditivo_service.listar_por_obra(obra_id)
        lancamentos = self._lancamento_service.listar_por_obra(obra_id)
        anexos = self._anexo_service.listar_por_obra(obra_id)
        
        # Validar tipos
        for lanc in lancamentos:
            if not getattr(lanc, "tipo_nome", None) or not lanc.tipo_nome.strip():
                logger.warning(
                    "Lançamento %s (%s) sem tipo associado",
                    getattr(lanc, "id", "?"),
                    lanc.descricao
                )
        
        # Criar documento DOCX
        doc = self._criar_documento()
        
        # === CABEÇALHO ===
        self._adicionar_titulo(doc, "RELATÓRIO DA OBRA", tamanho=14)
        self._adicionar_titulo(doc, _texto(obra.nome, "Obra sem nome"), tamanho=13, sublinhado=True)
        
        # Informações da obra em tabela invisível
        info_table = doc.add_table(rows=2, cols=2)
        info_table.style = None
        info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        info_table.cell(0, 0).text = f"Código: {_texto(obra.codigo)}"
        info_table.cell(0, 1).text = f"Cliente: {_texto(obra.cliente_contratante, 'Não informado')}"
        info_table.cell(1, 0).text = f"Local: {_texto(obra.local_obra, 'Não informado')}"
        info_table.cell(1, 1).text = f"Engenheiro: {_texto(obra.engenheiro_responsavel, 'Não informado')}"
        
        for row in info_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(9)
        
        # Emissão
        p_emissao = doc.add_paragraph()
        p_emissao.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_emissao = p_emissao.add_run(f"Emissão: {datetime.now().strftime('%d/%m/%Y')}")
        run_emissao.font.name = "Arial"
        run_emissao.font.size = Pt(9)
        
        doc.add_paragraph()
        
        # === RESUMO FINANCEIRO ===
        self._adicionar_secao(doc, "RESUMO FINANCEIRO")
        self._adicionar_tabela_resumo(doc, resumo)
        
        # === ADITIVOS ===
        if aditivos:
            self._adicionar_secao(doc, "ADITIVOS")
            self._adicionar_tabela_aditivos(doc, aditivos)
        
        # === LANÇAMENTOS ===
        if lancamentos:
            self._adicionar_secao(doc, "LANÇAMENTOS")
            self._adicionar_tabela_lancamentos(doc, lancamentos)
        
        # === ANEXOS ===
        if anexos:
            self._adicionar_secao(doc, "ANEXOS")
            self._adicionar_anexos(doc, anexos)
        
        # === ASSINATURA ===
        responsavel = self._obter_responsavel()
        cnpj = self._obter_cnpj()
        self._adicionar_assinatura(doc, responsavel, cnpj)
        
        # Salvar DOCX temporário
        filename = f"relatorio_obra_{obra.codigo}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        docx_path = self._storage.relatorio_path(filename)
        doc.save(str(docx_path))
        print(f"[DOCX] DOCX: {docx_path}")
        
        # Converter para PDF
        pdf_path = self._converter_para_pdf(docx_path)
        print(f"[DOCX] PDF: {pdf_path}")
        
        # Registrar no repositório
        from controle_obras.domain.models import RelatorioGerado
        
        self._relatorio_repo.save(
            RelatorioGerado(
                obra_id=obra_id,
                tipo_relatorio="obra_docx",
                arquivo_gerado=str(pdf_path),
            )
        )
        
        print(f"[DOCX] Status: sucesso")
        
        return pdf_path
